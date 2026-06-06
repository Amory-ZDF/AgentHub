"""
md_to_docx.py — 把 Markdown 转换为 Word (.docx)。
支持：标题（#~######）、段落、列表、表格（GFM）、代码块（```...```）、行内代码、粗体/斜体、引用块。
不依赖 pandoc，仅 python-docx。
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CODE_FONT = "Menlo"
BODY_FONT_EN = "Calibri"
BODY_FONT_CN = "PingFang SC"


def set_cn_font(run, font_name=BODY_FONT_CN):
    run.font.name = BODY_FONT_EN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def add_inline(paragraph, text):
    """处理行内 **bold** *italic* `code` 简单解析。"""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_cn_font(r)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
            set_cn_font(r)
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = CODE_FONT
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
            set_cn_font(r)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_cn_font(r)


def add_code_block(doc, code_lines, lang=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F2")
    pPr.append(shd)

    text = "\n".join(code_lines)
    r = p.add_run(text)
    r.font.name = CODE_FONT
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table(doc, header_cells, body_rows):
    cols = len(header_cells)
    table = doc.add_table(rows=1 + len(body_rows), cols=cols)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header_cells):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h.strip())
        r.bold = True
        set_cn_font(r)
    for ridx, row in enumerate(body_rows, start=1):
        cells = table.rows[ridx].cells
        for i in range(cols):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            txt = row[i].strip() if i < len(row) else ""
            add_inline(p, txt)


def parse_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line):
    s = line.strip()
    if not s.startswith("|"):
        return False
    cells = parse_table_row(s)
    return all(re.fullmatch(r":?-+:?", c) for c in cells if c)


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT_EN
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.lstrip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip()[3:].strip()
                code_buf = []
            else:
                add_code_block(doc, code_buf, code_lang)
                in_code = False
                code_buf = []
                code_lang = ""
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.rstrip()

        # 空行
        if not stripped.strip():
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading = doc.add_heading(level=min(level, 4))
            r = heading.add_run(text)
            set_cn_font(r)
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            r = p.add_run(text)
            r.italic = True
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            set_cn_font(r)
            i += 1
            continue

        # 表格
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = parse_table_row(stripped)
            body = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                body.append(parse_table_row(lines[j]))
                j += 1
            add_table(doc, header, body)
            i = j
            continue

        # 列表
        m = re.match(r"^(\s*)([-*])\s+(.*)$", stripped)
        if m:
            text = m.group(3)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text)
            i += 1
            continue
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", stripped)
        if m:
            text = m.group(3)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, text)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    if in_code and code_buf:
        add_code_block(doc, code_buf, code_lang)

    doc.save(str(docx_path))
    print(f"OK -> {docx_path}")


if __name__ == "__main__":
    pairs = [
        ("PRD-product.md", "PRD-product.docx"),
        ("PRD-engineering.md", "PRD-engineering.docx"),
    ]
    for src, dst in pairs:
        src_p = Path(src)
        if not src_p.exists():
            print(f"SKIP missing: {src}")
            continue
        md_to_docx(src_p, Path(dst))
